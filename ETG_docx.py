from docx import Document, enum
from docx.enum import text
from docx.shared import Inches, Pt
import os

#####################
# Margin Limits
UPPER_LIMIT_TB = 10.99  # 10.82
UPPER_LIMIT_LR = 7.83
LOWER_LIMIT_TBLR = 0.0  # 0.17
SPAN_LIMIT_TB = 10.99
SPAN_LIMIT_LR = 8
#####################

class doc_margins:
    def __init__(self, margins, tab_stops):
        self.margins = margins
        self.tab_stops = tab_stops

def create_doc(tag_directory, doc_title, **kwargs):
    # If there's no doc title, default to Operator Tags
    # If the extension passed through isn't .docx, rewrite it to that
    # If the doc title has no extension, add .docx
    if doc_title == "":
        doc_title = 'Operator Tags.docx'
    elif "." in doc_title:
        split_doc_title = doc_title.split(".")
        if split_doc_title[1] != ".docx":  # Rewrite file extension to .docx
            doc_title = split_doc_title[0] + ".docx"
    else:
        doc_title += ".docx"

    # Check that the file can be opened
    try:
        open(doc_title, "w")
    except IOError:
        print("Write access DENIED on %s" % doc_title)
        return False

    directory = os.fsencode(tag_directory)

    document = Document()

    # if nothing is provided, kwargs defaults to the second value, allowing no minimum
    # It acts like an absolute value where a margin of -2" is the same as 2"
    # The max value for top and bottom margin 10.82, the min value is 0.17
    # For a 11" tall sheet, the margins need a 0.01 difference minimum
    top_margin = round(max(min(UPPER_LIMIT_TB, abs(kwargs.get('top_margin', 0.5))), LOWER_LIMIT_TBLR), 2)
    bottom_margin = round(max(min(UPPER_LIMIT_TB, abs(kwargs.get('bot_margin', 0.5))), LOWER_LIMIT_TBLR), 2)
    if (top_margin + bottom_margin) > SPAN_LIMIT_TB: # Prioritize the top margin, adjust the bottom margin to conform
        bottom_margin = SPAN_LIMIT_TB - top_margin
    # The max value for left and right margin is 7.83, the min value is 0.17
    # For a 8.5" wide sheet, the margins need a 0.5 difference minimum
    left_margin = round(max(min(UPPER_LIMIT_LR, abs(kwargs.get('left_margin', 0.2))), LOWER_LIMIT_TBLR), 2)
    right_margin = round(max(min(UPPER_LIMIT_LR, abs(kwargs.get('right_margin', 0.2))), LOWER_LIMIT_TBLR), 2)
    if (left_margin + right_margin) > SPAN_LIMIT_LR: # Prioritize the left margin, adjust the right margin to conform
        right_margin = SPAN_LIMIT_LR - left_margin

    sections = document.sections
    for section in sections:
        # I had the top and bottom margins at 0.2 but there's a header and footer on the sheet that it's printed on
        section.top_margin = Inches(top_margin) # Inches(0.5)
        section.bottom_margin = Inches(bottom_margin) # Inches(0.4)
        section.left_margin = Inches(left_margin) # Inches(0.0) # 0.2 still leaves ~13/32"(0.40625") margin on either side, changing to 0.0"
        section.right_margin = Inches(right_margin) # Inches(0.0)

    p = document.add_paragraph()
    # # paragraph alignment is left by default
    # # Don't need to worry about alignment since it's using tab stops below to position the images
    # # p.alignment = 0 # for left, 1 for center, 2 = right, 3 = justify
    p.paragraph_format.line_spacing = 1.0
    space_after = Pt(kwargs.get('space_after', 0.0))
    p.paragraph_format.space_after = space_after

    # Word says the image is 2.63" wide but the real life print is almost exactly 2.5. That's over a 1/8" difference
    # The sticker that the image needs to fit on is 2 5/8"(2.625")
    # margin_offset is for the already implemented tab_stop left and tab_stop right. It might make more sense to change
    # it to the center of where the image is supposed to be and do a tab_stop center for all the images

    # When I had the left margin at 0.2", the printed image ended up at 0.25", so for a scalar, if I want to 0.2" margin
    # then I should use the scalar 0.2/0.25 to pass in a lower value margin to end up with a print at 0.2"
    # The edge of the sticker is 3/16"(0.1875") away from the edge of the page, so I want the image to be at that edge

    # The word doc shows the picture to be 2.63" but the printed image is 2 19/32"(2.59375") which is a little over
    # 1/32" difference. meaning that the above difference of the margin from 0.2" to 0.25" is a difference, not a scalar
    # I think if I were to just take the difference of the set and measured margin, I'll get a calculated margin that
    # will put the image within 1/32" of the edge of the sticker.
    margin_offset = (3 / 16) - (0.25 - 0.2) # This equals 0.1375" # (2.625 - 2.5) / 2

    margin_start = section.left_margin.inches + margin_offset # Inches(section.left_margin.inches)
    margin_middle = (section.page_width.inches - (section.left_margin.inches + section.right_margin.inches)) / 2 # Inches(section.page_width.inches / 2)
    margin_end = section.page_width.inches - (section.left_margin.inches + section.right_margin.inches) - margin_offset

    # Adding this here rather than at the top of the method so I can reuse margin_start, middle, and end
    tab_alignment = {
        'left': enum.text.WD_TAB_ALIGNMENT.LEFT,
        'center': enum.text.WD_TAB_ALIGNMENT.CENTER,
        'right': enum.text.WD_TAB_ALIGNMENT.RIGHT,
        'decimal': enum.text.WD_TAB_ALIGNMENT.DECIMAL,
        'bar': enum.text.WD_TAB_ALIGNMENT.BAR,
        'list': enum.text.WD_TAB_ALIGNMENT.LIST,
        'clear': enum.text.WD_TAB_ALIGNMENT.CLEAR,
        'end': enum.text.WD_TAB_ALIGNMENT.END,
        'num': enum.text.WD_TAB_ALIGNMENT.NUM,
        'start': enum.text.WD_TAB_ALIGNMENT.START
    }

    tab_entry = kwargs.get('tab', [])# [['left', margin_start], ['center', margin_middle], ['right', margin_end]])

    picture_height = kwargs.get('picture_height', 1.03125)
    picture_width = kwargs.get('picture_width', 2.63)

    r = p.add_run()

    tab_stops = p.paragraph_format.tab_stops
    for tab in tab_entry:
        # print(tab)
        tab_stops.add_tab_stop(Inches(tab[1]), tab_alignment[tab[0]])

    # if margin_start != 0: r.add_text('\t')  # Needed if using margin_start and it's non-zero
    # if len(tab_entry) != 0: r.add_text('\t')

    for index, file in enumerate(os.listdir(directory)):
        filename = os.fsdecode(file)
        if filename.endswith(".png"): #("_tag.png"):
            # # changed adding a picture with specified width to specified height, increased the height from 1" to 1 1/32"
            # r.add_picture(f'{tag_directory}\\{filename}', width=Inches(picture_width), height=Inches(picture_height))# width=Inches(2.63))
            # # print(f'{index} <> {directory_in_str}\\{filename} +++ Added')
            # # After 3 images are added, add new paragraph to carry over tab_stops
            if len(tab_entry) == 0:
                pass
            elif (index) % len(tab_entry) == 0: # 3 == 0: # after 3 images are loaded, start a new paragraph
                # print(filename)
                p = document.add_paragraph() # need to add a paragraph to start a new line
                # Each time a new paragraph is added, the same parameters need to be configured
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = space_after

                tab_stops = p.paragraph_format.tab_stops
                for tab in tab_entry:
                    # print(tab)
                    tab_stops.add_tab_stop(Inches(tab[1]), tab_alignment[tab[0]])
                r = p.add_run()
                # if margin_start != 0: r.add_text('\t')  # Needed if using margin_start and it's non-zero
                r.add_text('\t')
                # r.add_text('\n')
            else:
                r.add_text('\t') # tab
                # r.add_text('  ')

            r.add_picture(f'{tag_directory}\\{filename}', width=Inches(picture_width), height=Inches(picture_height))

    document.save(doc_title)
    return True

if __name__ == "__main__":
    input_tab = [('left', 0.0), ('left', 0.7), ('left', 1.3), ('left', 1.9), ('left', 2.5), ('left', 3.1), ('left', 3.7)]
    create_doc("C:\\pythonProject\\Employee_Tag_Generator\\generated_images\\id_images",
               "Operator Tags Test V2", top_margin=10.99, bot_margin=0.0, left_margin=0.0, right_margin=0.0,
               space_after=0.0, picture_width=0.5, picture_height=0.5)